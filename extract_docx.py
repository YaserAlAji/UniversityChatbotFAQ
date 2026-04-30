from docx import Document
import os
import sys

# Set stdout to use utf-8 to avoid encoding errors in windows console
sys.stdout.reconfigure(encoding='utf-8')

filename = 'Ai Chatbot University Project.docx'
if not os.path.exists(filename):
    print(f"File not found: {filename}")
    exit(1)

try:
    doc = Document(filename)

    print("--- FILE CONTENT START ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)

    print("\n--- TABLES START ---")
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            print(" | ".join(row_text))
        print("--- TABLE END ---")

    print("--- FILE CONTENT END ---")
except Exception as e:
    print(f"Error reading file: {e}")
